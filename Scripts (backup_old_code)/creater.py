import tools as t
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx as docx
import random
import bestaetigungsfenster as b

doc = docx.Document()

class Creater:
    def __init__(self, quelldatei, ges):
        self.vokabelanzahl = 0
        self.quelldatei = quelldatei
        self.wortschatz = f'{quelldatei[quelldatei.rfind("/")+1:quelldatei.rfind(".")]}'
        self.ges = ges

        self.lektionsangabe = ""
        for lektion in self.ges:
            self.lektionsangabe += f"{lektion}+"
        self.lektionsangabe = self.lektionsangabe.strip("+")

        for lektion in self.ges:
            self.vokabelanzahl += int(self.ges[lektion][0].split(":")[1])

        self.doc = doc
        self.datum = t.datumsanzeige()

    def text(self, text, schriftgroesse, schriftstil, *extras):
        para = self.doc.add_paragraph()

        if "RIGHT" in extras:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if "CENTER" in extras:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "DISTRIBUTE" in extras:
           para.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE

        para = para.add_run(text)
        para.font.size = Pt(schriftgroesse)
        para.font.name = schriftstil

        if "underline" in extras:
            para.underline = True
        if "bold" in extras:
                para.bold = True

    def create_line(self):
        space = ""
        for i in range(100):
            space += " "
        line = space + "_"
        return line

    def create_test(self):
        self.anordnung_gruppe = {"A": [], "B": []}

        #print(f"{self.ges}\n")

        for lektion in self.ges:
            #choosing the vocabs
            t.vokabelauswahl("subs", lektion, self.ges, self.quelldatei, self.anordnung_gruppe)
            t.vokabelauswahl("verb", lektion, self.ges, self.quelldatei, self.anordnung_gruppe)
            t.vokabelauswahl("adje", lektion, self.ges, self.quelldatei, self.anordnung_gruppe)
            t.vokabelauswahl("klwo", lektion, self.ges, self.quelldatei, self.anordnung_gruppe)

        #mixing the order of the vocabs
        for i in range(len(self.anordnung_gruppe["A"])):
            if self.anordnung_gruppe["A"][i] == self.anordnung_gruppe["B"][i]:
                index = random.randint(0, len(self.anordnung_gruppe["A"]) - 1)
                while index == i:
                    index = random.randint(0, len(self.anordnung_gruppe["A"]) - 1)
                self.anordnung_gruppe["A"][i], self.anordnung_gruppe["A"][index] = self.anordnung_gruppe["A"][index], self.anordnung_gruppe["A"][i]
        #print(f'Test A: {self.anordnung_gruppe["A"]}\nTest B: {self.anordnung_gruppe["B"]}')

        file = open(t.find_path(f"Output\Vokabeltest_{self.wortschatz}_L{self.lektionsangabe}_{self.datum}.txt"), "w")
        for i in self.anordnung_gruppe:
            file.write(f"Test {i}:\n\n")
            self.text(self.datum, 9, "Calibri", "RIGHT")
            self.text(f"({i}) Vokabeltest - Lektion {self.lektionsangabe}", 9, "Calibri", "underline", "bold", "CENTER")
            self.text("Übersetzen Sie die folgenden Vokabeln und ergänzen Sie bei Nomen Genitiv Sg. und Geschlecht, bei den Verben\n""Stammformen, bei Präpositionen den Kasus, mit dem sie stehen, und bei Adjektiven und Pronomen die Genera.",9, "Calibri")
            for k in range(self.vokabelanzahl):
                vokabel = self.anordnung_gruppe[i][k]
                if vokabel.rstrip().endswith(")"):
                    vokabel = vokabel.split('(')[0]
                self.text(f"{vokabel} {self.create_line()}", 9,"Calibri", "underline", "DISTRIBUTE")
                file.write(f"{vokabel}\n")
            file.write("\n\n")
        file.close() 

        self.doc.save(t.find_path(f"Output\Vokabeltest_{self.wortschatz}_L{self.lektionsangabe}_{self.datum}.docx"))
        b.Bestaetigungsfenster(self.wortschatz, self.lektionsangabe, self.datum).start() 


