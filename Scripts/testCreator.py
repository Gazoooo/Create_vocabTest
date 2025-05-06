import tools as t
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx as docx

#creates the .docx and .txt files for the test
class TestCreator:
    """
    Creates .docx and .txt files for a vocab test based on the provided vocabulary data.
    """
    
    def __init__(self, wortschatz, all_data, anordnung_gruppe):
        """
        Initializes the TestCreator with the given vocabulary set, lesson data, and group arrangement.

        Args:
            wortschatz (str): The vocabulary set name.
            all_data (dict): A dictionary containing lesson data with the lesson name as key.
            anordnung_gruppe (dict): A dictionary with the vocabulary arrangement for each group.
        """
        self.datum = t.datumsanzeige()
        self.wortschatz = wortschatz
        self.lektionsangabe = ""
        self.vokabelanzahl = 0
        for lektion in all_data:
            self.lektionsangabe += f"{lektion}+"
            self.vokabelanzahl += int(all_data[lektion][0].split(":")[1]) #used not for filename, but later
        self.lektionsangabe = self.lektionsangabe.strip("+")
        self.filename = f"Vokabeltest_{self.wortschatz}_L{self.lektionsangabe}_{self.datum}"

        self.anordnung_gruppe = anordnung_gruppe

    def create_docx(self):
        """
        Creates a .docx file containing the vocabulary test with proper formatting and saves it 
        under the `"Output/Vokabeltest/"` directory.

        The method adds lesson details, instructions, and the vocabularies (with blank lines for 
        translations) for each group in the test.
        """
        self.doc = docx.Document()
        for gruppe in self.anordnung_gruppe:
            self.text(self.datum, 9, "Calibri", "RIGHT")
            self.text(f"({gruppe}) Vokabeltest - Lektion {self.lektionsangabe}", 9, "Calibri", "underline", "bold", "CENTER")
            self.text("Übersetzen Sie die folgenden Vokabeln und ergänzen Sie bei Nomen Genitiv Sg. und Geschlecht, bei den Verben\n""Stammformen, bei Präpositionen den Kasus, mit dem sie stehen, und bei Adjektiven und Pronomen die Genera.",9, "Calibri")
            for voc in range(self.vokabelanzahl):
                vokabel = self.anordnung_gruppe[gruppe][voc]
                if vokabel.rstrip().endswith(")"):
                    vokabel = vokabel.split('(')[0]
                self.text(f"{vokabel} {self.create_line()}", 9,"Calibri", "underline", "DISTRIBUTE")
        self.doc.save(t.find_path(f"Output/{self.filename}.docx"))

    def create_txt(self):
        """
        Creates a .txt file containing the vocabulary test with plain text content and saves it 
        under the `"Output/Vokabeltest/"` directory.

        The method writes each group's vocabulary list in the test, without any formatting or blank lines.
        """
        file = open(t.find_path(f"Output/{self.filename}.txt"), "w")
        for gruppe in self.anordnung_gruppe:
            file.write(f"Test {gruppe}:\n\n")
            for voc in range(self.vokabelanzahl):
                vokabel = self.anordnung_gruppe[gruppe][voc]
                if vokabel.rstrip().endswith(")"):
                    vokabel = vokabel.split('(')[0]
                file.write(f"{vokabel}\n")
            file.write("\n\n")
        file.close() 

    def text(self, text, schriftgroesse, schriftstil, *extras):
        """
        A convenience method to add a paragraph of text to the .docx file with specified formatting.

        Args:
            text (str): The text to be added to the paragraph.
            schriftgroesse (int): The font size of the text.
            schriftstil (str): The font style (e.g., "Calibri").
            *extras (str): Optional additional formatting parameters (e.g., "underline", "bold", "CENTER").
        """
        para = self.doc.add_paragraph()
        if "RIGHT" in extras:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if "CENTER" in extras:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "DISTRIBUTE" in extras:
           para.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        para = para.add_run(text)
        para.font.size = Pt(schriftgroesse)
        para.font.name = schriftstil
        if "underline" in extras:
            para.underline = True
        if "bold" in extras:
                para.bold = True

    def create_line(self):
        """
        Creates a line of underscores to be used for leaving space for the answers in the vocab test.

        Returns:
            str: A string containing 100 spaces followed by an underscore ("_").
        """
        space = ""
        for i in range(100):
            space += " "
        line = space + "_"
        return line